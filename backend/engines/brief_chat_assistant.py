"""
Brief Chat Assistant - Conversational AI for discussing generated briefs

Uses Groq API (Llama 3.1) for fast, context-aware conversations about
procurement briefs. Automatically loads brief content when DOCX files
are generated.

Features:
- Full conversation history
- Brief-aware context (knows the generated documents)
- Procurement-focused responses
- Fast responses via Groq
"""

import os
import requests
from typing import Dict, Any, List, Optional
from pathlib import Path
from docx import Document


class BriefChatAssistant:
    """
    Conversational assistant for discussing generated procurement briefs.

    Uses Groq API with Llama 3.1 for fast, intelligent responses.
    Maintains conversation history and brief context.
    """

    GROQ_API_URL = "https://api.groq.com/openai/v1/chat/completions"
    MODEL = "llama-3.3-70b-versatile"  # Latest fast model

    def __init__(self, api_key: Optional[str] = None):
        """
        Initialize chat assistant with Groq API.

        Args:
            api_key: Groq API key (if None, reads from environment)
        """
        self.api_key = api_key or os.getenv('GROQ_API_KEY', '')
        self.enabled = bool(self.api_key)

        # Conversation state
        self.conversation_history: List[Dict[str, str]] = []
        self.brief_context: Dict[str, Any] = {}
        self.subcategory: str = ""

        if not self.enabled:
            print("[WARN] GROQ_API_KEY not set - chat assistant disabled")
        else:
            print("[OK] Brief Chat Assistant initialized with Groq API")

    def load_brief_context(
        self,
        incumbent_path: Optional[str] = None,
        regional_path: Optional[str] = None,
        subcategory: str = "",
        brief_data: Optional[Dict[str, Any]] = None
    ):
        """
        Load generated brief content for context-aware conversations.

        Called automatically when DOCX files are generated.

        Args:
            incumbent_path: Path to incumbent concentration brief DOCX
            regional_path: Path to regional concentration brief DOCX
            subcategory: The procurement subcategory
            brief_data: Optional raw brief data dict
        """
        self.subcategory = subcategory
        self.brief_context = {
            'subcategory': subcategory,
            'incumbent_content': '',
            'regional_content': '',
            'summary': ''
        }

        # Extract text from incumbent brief
        if incumbent_path and Path(incumbent_path).exists():
            self.brief_context['incumbent_content'] = self._extract_docx_text(incumbent_path)
            self.brief_context['incumbent_path'] = incumbent_path

        # Extract text from regional brief
        if regional_path and Path(regional_path).exists():
            self.brief_context['regional_content'] = self._extract_docx_text(regional_path)
            self.brief_context['regional_path'] = regional_path

        # Store raw data if provided
        if brief_data:
            self.brief_context['raw_data'] = brief_data

        # Create summary for system prompt
        self._create_brief_summary()

        # Reset conversation with new context
        self.conversation_history = []

        print(f"[OK] Loaded brief context for: {subcategory}")

    def _extract_docx_text(self, docx_path: str, max_chars: int = 8000) -> str:
        """Extract text content from DOCX file."""
        try:
            doc = Document(docx_path)
            paragraphs = []

            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    paragraphs.append(text)

            # Also extract from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = ' | '.join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        paragraphs.append(row_text)

            full_text = '\n'.join(paragraphs)

            # Truncate if too long
            if len(full_text) > max_chars:
                full_text = full_text[:max_chars] + "\n[... content truncated ...]"

            return full_text

        except Exception as e:
            print(f"[WARN] Failed to extract DOCX text: {e}")
            return ""

    def _create_brief_summary(self):
        """Create a summary of the brief content for the system prompt."""
        inc = self.brief_context.get('incumbent_content', '')
        reg = self.brief_context.get('regional_content', '')

        summary_parts = []

        if inc:
            # Extract key sections
            summary_parts.append(f"INCUMBENT CONCENTRATION BRIEF:\n{inc[:3000]}")

        if reg:
            summary_parts.append(f"REGIONAL CONCENTRATION BRIEF:\n{reg[:3000]}")

        self.brief_context['summary'] = '\n\n'.join(summary_parts)

    def _get_system_prompt(self) -> str:
        """Build the system prompt with brief context."""
        base_prompt = """You are a procurement intelligence assistant helping analyze supplier briefs. You have full knowledge of the generated procurement briefs and can answer questions, provide recommendations, and discuss strategies.

IMPORTANT GUIDELINES:
1. Be concise but thorough in responses
2. Reference specific data from the briefs when answering
3. Provide actionable recommendations when asked
4. If asked about something not in the briefs, say so clearly
5. Use procurement industry terminology appropriately
6. Be conversational and helpful

"""

        if self.brief_context.get('summary'):
            context_prompt = f"""BRIEF CONTEXT - You have access to these generated briefs for {self.subcategory}:

{self.brief_context['summary']}

When the user asks questions, use this brief data to provide accurate, grounded responses.
"""
            return base_prompt + context_prompt

        return base_prompt + "No briefs have been loaded yet. Ask the user to generate briefs first."

    def chat(self, user_message: str) -> Dict[str, Any]:
        """
        Send a message and get a response.

        Args:
            user_message: The user's message

        Returns:
            Dict with response text and metadata
        """
        if not self.enabled:
            return {
                'success': False,
                'response': "Chat assistant is not available. Please set GROQ_API_KEY in your .env file.",
                'error': 'API key not configured'
            }

        if not self.brief_context.get('summary'):
            return {
                'success': True,
                'response': "I don't have any brief context loaded yet. Please generate briefs first, then I can help you analyze them and provide recommendations.",
                'has_context': False
            }

        # Add user message to history
        self.conversation_history.append({
            'role': 'user',
            'content': user_message
        })

        try:
            # Build messages for API
            messages = [
                {'role': 'system', 'content': self._get_system_prompt()}
            ]

            # Add conversation history (last 10 messages to stay within limits)
            messages.extend(self.conversation_history[-10:])

            # Call Groq API
            headers = {
                'Authorization': f'Bearer {self.api_key}',
                'Content-Type': 'application/json'
            }

            payload = {
                'model': self.MODEL,
                'messages': messages,
                'temperature': 0.7,
                'max_tokens': 1024,
                'top_p': 0.9
            }

            response = requests.post(
                self.GROQ_API_URL,
                headers=headers,
                json=payload,
                timeout=30
            )
            response.raise_for_status()

            result = response.json()
            assistant_message = result.get('choices', [{}])[0].get('message', {}).get('content', '')

            # Add assistant response to history
            self.conversation_history.append({
                'role': 'assistant',
                'content': assistant_message
            })

            return {
                'success': True,
                'response': assistant_message,
                'has_context': True,
                'subcategory': self.subcategory,
                'model': self.MODEL
            }

        except requests.exceptions.Timeout:
            return {
                'success': False,
                'response': "Request timed out. Please try again.",
                'error': 'timeout'
            }
        except requests.exceptions.RequestException as e:
            error_msg = str(e)
            # Try to extract error message from response
            try:
                error_data = e.response.json()
                error_msg = error_data.get('error', {}).get('message', str(e))
            except:
                pass

            return {
                'success': False,
                'response': f"API error: {error_msg}",
                'error': error_msg
            }
        except Exception as e:
            return {
                'success': False,
                'response': f"Error: {str(e)}",
                'error': str(e)
            }

    def clear_history(self):
        """Clear conversation history but keep brief context."""
        self.conversation_history = []
        print("[OK] Conversation history cleared")

    def get_suggested_questions(self) -> List[str]:
        """Get suggested questions based on the loaded brief."""
        if not self.brief_context.get('summary'):
            return [
                "Generate briefs first to get suggestions",
            ]

        return [
            f"What are the main risks in our {self.subcategory} procurement?",
            "Which suppliers should we consider for diversification?",
            "What's the recommended action plan?",
            "Explain the supplier concentration risk",
            "What are the cost implications of switching suppliers?",
            "How can we reduce regional dependency?",
        ]

    def is_ready(self) -> bool:
        """Check if assistant is ready for chat."""
        return self.enabled and bool(self.brief_context.get('summary'))


# Convenience function
def create_chat_assistant() -> BriefChatAssistant:
    """Create a new chat assistant instance."""
    return BriefChatAssistant()
