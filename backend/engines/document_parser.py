"""
Document Parser - Handles uploads (PDF, Excel, Word, CSV)
Extracts spend data from various formats
"""

import pandas as pd
from pathlib import Path
from typing import Dict, Any, List, Tuple
from docx import Document as DocxDocument
import PyPDF2
import io


class DocumentParser:
    """Parse uploaded documents and extract spend data"""
    
    def __init__(self):
        self.supported_formats = ['.pdf', '.xlsx', '.xls', '.csv', '.docx']
    
    def parse_upload(self, file_content: bytes, filename: str) -> Dict[str, Any]:
        """
        Parse uploaded file and extract procurement data
        
        Args:
            file_content: File bytes
            filename: Original filename
            
        Returns:
            Dictionary with extracted spend data
        """
        file_ext = Path(filename).suffix.lower()
        
        if file_ext == '.csv':
            return self._parse_csv(file_content)
        elif file_ext == '.xlsx' or file_ext == '.xls':
            return self._parse_excel(file_content)
        elif file_ext == '.pdf':
            return self._parse_pdf(file_content)
        elif file_ext == '.docx':
            return self._parse_docx(file_content)
        else:
            return {'error': f'Unsupported format: {file_ext}'}
    
    def _parse_csv(self, content: bytes) -> Dict[str, Any]:
        """Parse CSV file"""
        try:
            df = pd.read_csv(io.BytesIO(content))
            return {
                'success': True,
                'data': df,
                'format': 'CSV',
                'rows': len(df),
                'columns': list(df.columns)
            }
        except Exception as e:
            return {'error': str(e), 'format': 'CSV'}
    
    def _parse_excel(self, content: bytes) -> Dict[str, Any]:
        """Parse Excel file"""
        try:
            df = pd.read_excel(io.BytesIO(content))
            return {
                'success': True,
                'data': df,
                'format': 'Excel',
                'rows': len(df),
                'columns': list(df.columns)
            }
        except Exception as e:
            return {'error': str(e), 'format': 'Excel'}
    
    def _parse_pdf(self, content: bytes) -> Dict[str, Any]:
        """Parse PDF file and extract tables"""
        try:
            reader = PyPDF2.PdfReader(io.BytesIO(content))
            text = ""
            for page in reader.pages:
                text += page.extract_text()
            
            return {
                'success': True,
                'data': text,
                'format': 'PDF',
                'pages': len(reader.pages),
                'note': 'Manual data extraction may be needed'
            }
        except Exception as e:
            return {'error': str(e), 'format': 'PDF'}
    
    def _parse_docx(self, content: bytes) -> Dict[str, Any]:
        """Parse DOCX file"""
        try:
            doc = DocxDocument(io.BytesIO(content))
            text = ""
            tables_data = []
            
            for para in doc.paragraphs:
                if para.text.strip():
                    text += para.text + "\n"
            
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text for cell in row.cells]
                    table_data.append(row_data)
                tables_data.append(table_data)
            
            return {
                'success': True,
                'text': text,
                'tables': tables_data,
                'format': 'DOCX',
                'tables_count': len(tables_data)
            }
        except Exception as e:
            return {'error': str(e), 'format': 'DOCX'}
    
    def extract_spend_data(self, df: pd.DataFrame) -> Dict[str, Any]:
        """
        Extract key procurement metrics from dataframe
        Works with various column naming conventions
        """
        columns_lower = [col.lower() for col in df.columns]
        
        result = {
            'total_rows': len(df),
            'columns': list(df.columns),
            'data': df,
            'metrics': {}
        }
        
        # Find spend column
        spend_col = None
        for col in columns_lower:
            if any(x in col for x in ['spend', 'cost', 'amount', 'value', 'usd']):
                spend_col = df.columns[columns_lower.index(col)]
                break
        
        if spend_col:
            try:
                total_spend = pd.to_numeric(df[spend_col], errors='coerce').sum()
                result['metrics']['total_spend'] = total_spend
            except:
                pass
        
        # Find category column
        category_col = None
        for col in columns_lower:
            if any(x in col for x in ['category', 'product', 'item', 'service']):
                category_col = df.columns[columns_lower.index(col)]
                break
        
        if category_col:
            result['metrics']['categories'] = df[category_col].unique().tolist()
        
        # Find supplier column
        supplier_col = None
        for col in columns_lower:
            if any(x in col for x in ['supplier', 'vendor', 'provider']):
                supplier_col = df.columns[columns_lower.index(col)]
                break
        
        if supplier_col:
            result['metrics']['suppliers'] = df[supplier_col].unique().tolist()
        
        return result
