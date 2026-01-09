"""
Enhanced DOCX & PDF Exporter
Converts leadership brief dictionaries to professional DOCX and PDF files
Matches exact format specification with additional insights

Features:
- Secure path handling with sanitization
- Automatic cleanup of failed exports
- Proper logging
"""

import os
import re
import logging
import tempfile
import shutil
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from typing import Dict, Any, List, Optional
from pathlib import Path
from datetime import datetime
import numpy as np

# Configure logger
logger = logging.getLogger(__name__)

try:
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4, letter
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch, cm
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
    from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False


class DOCXExportError(Exception):
    """Exception raised for DOCX export errors"""
    pass


class DOCXExporter:
    """
    Enhanced DOCX Exporter for Leadership Briefs
    - Matches exact format specification
    - Includes all new sections (Risk Matrix, ROI, Timeline, etc.)
    - Industry-agnostic
    - Organizes outputs into subfolders by category
    """

    def __init__(self, output_dir: str = "./outputs/briefs"):
        self.base_output_dir = Path(output_dir)
        self.base_output_dir.mkdir(parents=True, exist_ok=True)
        self.output_dir = self.base_output_dir  # Default, will be updated per export

        self.HEADER_COLOR = RGBColor(0, 51, 102)
        self.ACCENT_COLOR = RGBColor(0, 102, 153)

    @staticmethod
    def sanitize_filename(name: str, max_length: int = 100) -> str:
        """
        Securely sanitize a string for use in file/folder names.
        Prevents path traversal attacks and invalid characters.

        Args:
            name: The string to sanitize
            max_length: Maximum length of the result

        Returns:
            Sanitized string safe for filesystem use
        """
        if not name:
            return 'unknown'

        # First, normalize unicode and strip whitespace
        cleaned = name.strip()

        # Replace common special chars with readable alternatives
        replacements = {
            '&': 'and',
            '/': '_',
            '\\': '_',
            ':': '_',
            '*': '',
            '?': '',
            '"': '',
            '<': '',
            '>': '',
            '|': '_',
            ' ': '_',
            '.': '_',
            '..': '_',  # Prevent path traversal
        }

        for old, new in replacements.items():
            cleaned = cleaned.replace(old, new)

        # Remove any remaining non-alphanumeric characters except underscore and hyphen
        cleaned = re.sub(r'[^\w\-]', '', cleaned)

        # Collapse multiple underscores
        cleaned = re.sub(r'_+', '_', cleaned)

        # Remove leading/trailing underscores
        cleaned = cleaned.strip('_')

        # Truncate to max length
        if len(cleaned) > max_length:
            cleaned = cleaned[:max_length]

        # Ensure we have something
        return cleaned or 'unnamed'

    def _get_category_folder(self, brief_data: Dict[str, Any]) -> Path:
        """
        Create and return an organized subfolder path based on category hierarchy.
        Structure: outputs/briefs/{Sector}/{Category}/

        Uses secure path sanitization to prevent path traversal attacks.
        """
        # Extract hierarchy info
        sector = brief_data.get('sector', '')
        category = brief_data.get('category', 'General')

        # Build folder path with sanitized names
        folder_parts = []

        if sector:
            sanitized_sector = self.sanitize_filename(sector)
            if sanitized_sector:
                folder_parts.append(sanitized_sector)

        if category:
            sanitized_category = self.sanitize_filename(category)
            if sanitized_category:
                folder_parts.append(sanitized_category)

        if folder_parts:
            # Use Path properly to avoid path traversal
            subfolder = self.base_output_dir
            for part in folder_parts:
                subfolder = subfolder / part
        else:
            subfolder = self.base_output_dir / 'General'

        # Verify the path is within base_output_dir (prevent path traversal)
        try:
            subfolder = subfolder.resolve()
            base_resolved = self.base_output_dir.resolve()
            if not str(subfolder).startswith(str(base_resolved)):
                logger.warning(f"Path traversal attempt detected: {subfolder}")
                subfolder = self.base_output_dir / 'General'
        except (OSError, ValueError) as e:
            logger.warning(f"Path resolution error: {e}")
            subfolder = self.base_output_dir / 'General'

        # Create the folder
        try:
            subfolder.mkdir(parents=True, exist_ok=True)
        except OSError as e:
            logger.error(f"Failed to create output folder: {e}")
            # Fall back to base directory
            subfolder = self.base_output_dir
            subfolder.mkdir(parents=True, exist_ok=True)

        return subfolder
    
    def _set_margins(self, doc: Document):
        for section in doc.sections:
            section.top_margin = Inches(0.75)
            section.bottom_margin = Inches(0.75)
            section.left_margin = Inches(0.75)
            section.right_margin = Inches(0.75)
    
    def _add_title_section(self, doc: Document, title: str, subtitle: str, total_spend: float, fiscal_year: int):
        title_para = doc.add_paragraph()
        title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        title_run = title_para.add_run(title)
        title_run.font.size = Pt(16)
        title_run.font.bold = True
        title_run.font.color.rgb = self.HEADER_COLOR
        
        if subtitle:
            subtitle_para = doc.add_paragraph()
            subtitle_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            subtitle_run = subtitle_para.add_run(subtitle)
            subtitle_run.font.size = Pt(11)
            subtitle_run.font.italic = True
        
        spend_para = doc.add_paragraph()
        spend_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        spend_run = spend_para.add_run(
            f"Total Category Spend {fiscal_year}: USD {float(total_spend):,.0f}"
        )
        spend_run.font.size = Pt(12)
        spend_run.font.bold = True
        
        doc.add_paragraph()
    
    def _convert_numpy(self, value):
        if isinstance(value, np.integer):
            return int(value)
        elif isinstance(value, np.floating):
            return float(value)
        return value
    
    def _add_section_heading(self, doc: Document, text: str, level: int = 2):
        heading = doc.add_heading(text, level=level)
        for run in heading.runs:
            run.font.color.rgb = self.HEADER_COLOR
    
    def _create_styled_table(self, doc: Document, headers: List[str], rows: List[List[str]]):
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = 'Table Grid'
        
        header_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            header_cells[i].text = header
            for paragraph in header_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.size = Pt(10)
        
        for row_data in rows:
            row_cells = table.add_row().cells
            for i, cell_data in enumerate(row_data):
                row_cells[i].text = str(cell_data)
                for paragraph in row_cells[i].paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(10)
        
        return table
    
    def export_incumbent_concentration_brief(
        self, 
        brief_data: Dict[str, Any],
        filename: str = None
    ) -> str:
        """Export incumbent concentration brief with all enhanced sections"""
        doc = Document()
        self._set_margins(doc)
        
        category = brief_data.get('category', 'Procurement')
        total_spend = float(brief_data.get('total_spend', 0))
        fiscal_year = brief_data.get('fiscal_year', datetime.now().year)
        
        self._add_title_section(
            doc,
            brief_data.get('title', f'LEADERSHIP BRIEF â€“ {category.upper()} PROCUREMENT DIVERSIFICATION'),
            brief_data.get('subtitle', 'Supplier Concentration Analysis & Diversification Strategy'),
            total_spend,
            fiscal_year
        )
        
        current_state = brief_data.get('current_state', {})
        dominant_supplier = current_state.get('dominant_supplier', 'Primary Supplier')
        supplier_location = current_state.get('supplier_location', 'Multiple Regions')
        spend_share_pct = self._convert_numpy(current_state.get('spend_share_pct', 90))
        spend_share_usd = self._convert_numpy(current_state.get('spend_share_usd', 0))
        alternate_supplier = current_state.get('alternate_supplier_active', 'None active today')
        key_risk = current_state.get('key_risk', 'Supply chain risk')
        
        self._add_section_heading(doc, 'CURRENT STATE â€“ SUPPLIER CONCENTRATION RISK')
        
        metrics_table = self._create_styled_table(
            doc,
            ['Metric', f'{fiscal_year} Original'],
            [
                ['Supplier', dominant_supplier],
                ['Supplier Location', supplier_location],
                ['Spend Share', f'{spend_share_pct:.0f}% = USD {spend_share_usd:,.0f}'],
                [f'Current {category} Buying from Supplier', 'Yes'],
                [f'Alternate Supplier for {category}', alternate_supplier],
                ['Key Risk', key_risk]
            ]
        )
        
        doc.add_paragraph()
        
        self._add_section_heading(doc, 'OVERALL RISK STATEMENT', level=3)
        
        risk_statement = brief_data.get('risk_statement', '')
        if risk_statement:
            para = doc.add_paragraph(risk_statement)
            para.paragraph_format.space_after = Pt(12)
        
        doc.add_paragraph()
        
        self._add_section_heading(doc, 'SUPPLIER SHARE REDUCTION IMPACT (TARGET STATE)')
        
        supplier_reduction = brief_data.get('supplier_reduction', {})
        dominant_info = supplier_reduction.get('dominant_supplier', {})
        alternate_info = supplier_reduction.get('alternate_supplier', {})
        
        original_pct = self._convert_numpy(dominant_info.get('original_share_pct', spend_share_pct))
        new_pct = self._convert_numpy(dominant_info.get('new_target_cap_pct', 62))
        reduction_formatted = dominant_info.get('formatted_reduction', f'{original_pct:.0f}% â†’ {new_pct:.0f}%')
        
        alternate_new_pct = self._convert_numpy(alternate_info.get('new_target_pct', 38))
        
        self._create_styled_table(
            doc,
            ['Supplier', 'Original Share', 'New Target Cap', 'Reduction in Concentration'],
            [
                [dominant_supplier, f'{original_pct:.0f}%', f'{new_pct:.0f}%', reduction_formatted],
                [alternate_info.get('name', 'Alternate Supplier'), '0%', f'{alternate_new_pct:.0f}%', 'Enables supplier competition + fallback']
            ]
        )
        
        doc.add_paragraph()
        
        self._add_section_heading(doc, 'REGIONAL DEPENDENCY IMPROVEMENT')
        
        regional_dep = brief_data.get('regional_dependency', {})
        corridor_name = regional_dep.get('corridor_name', 'Primary Supply Corridor')
        original_region_pct = self._convert_numpy(regional_dep.get('original_pct', 90))
        new_target_pct = regional_dep.get('new_target_pct', '58â€“63%')
        net_reduction = regional_dep.get('net_reduction_pct', '27â€“32%')
        
        self._create_styled_table(
            doc,
            ['Metric', f'Original {fiscal_year}', 'New Target', 'Net Reduction'],
            [
                [corridor_name, f'{original_region_pct:.0f}%', str(new_target_pct), str(net_reduction)]
            ]
        )
        
        doc.add_paragraph()
        
        alt_supplier_name = alternate_info.get('name', 'Alternate Supplier')
        self._add_section_heading(doc, f'COST ADVANTAGE PROJECTION FROM ADDING {alt_supplier_name.upper()} INTO {category.upper()} PROCUREMENT')
        
        cost_advantages = brief_data.get('cost_advantages', [])
        cost_rows = []
        for advantage in cost_advantages:
            if isinstance(advantage, dict):
                region = advantage.get('region', '')
                driver = advantage.get('driver', '')
                min_usd = self._convert_numpy(advantage.get('min_usd', 0))
                max_usd = self._convert_numpy(advantage.get('max_usd', 0))
                cost_rows.append([region, driver, f'${min_usd:,.0f} â€“ ${max_usd:,.0f}'])
        
        if cost_rows:
            self._create_styled_table(
                doc,
                [f'Region of {alt_supplier_name} Presence', 'Advantage Driver', 'Estimated Annual Cost Benefit (USD)'],
                cost_rows
            )
        
        doc.add_paragraph()
        
        risk_matrix = brief_data.get('risk_matrix', {})
        if risk_matrix:
            self._add_section_heading(doc, 'RISK ASSESSMENT MATRIX')
            
            self._create_styled_table(
                doc,
                ['Risk Category', 'Current Level', 'Description'],
                [
                    ['Supply Chain Risk', risk_matrix.get('supply_chain_risk', 'N/A'), 'Single supplier dependency level'],
                    ['Geographic Risk', risk_matrix.get('geographic_risk', 'N/A'), 'Regional concentration level'],
                    ['Supplier Diversity Risk', risk_matrix.get('supplier_diversity_risk', 'N/A'), 'Number of qualified suppliers'],
                    ['Overall Risk Score', f"{risk_matrix.get('risk_score', 0)}/4.0", risk_matrix.get('overall_risk', 'N/A')]
                ]
            )
            doc.add_paragraph()
        
        roi = brief_data.get('roi_projections', {})
        if roi:
            self._add_section_heading(doc, 'ROI PROJECTIONS')
            
            annual_savings_min = self._convert_numpy(roi.get('annual_cost_savings_min', 0))
            annual_savings_max = self._convert_numpy(roi.get('annual_cost_savings_max', 0))
            risk_value = self._convert_numpy(roi.get('risk_reduction_value', 0))
            impl_cost = self._convert_numpy(roi.get('implementation_cost', 0))
            roi_min = roi.get('roi_percentage_min', 0)
            roi_max = roi.get('roi_percentage_max', 0)
            payback_min = roi.get('payback_period_months_min', 0)
            payback_max = roi.get('payback_period_months_max', 0)
            three_year_min = self._convert_numpy(roi.get('three_year_net_benefit_min', 0))
            three_year_max = self._convert_numpy(roi.get('three_year_net_benefit_max', 0))
            
            self._create_styled_table(
                doc,
                ['Financial Metric', 'Projected Value'],
                [
                    ['Annual Cost Savings', f'${annual_savings_min:,.0f} â€“ ${annual_savings_max:,.0f}'],
                    ['Risk Reduction Value', f'${risk_value:,.0f}'],
                    ['Implementation Cost', f'${impl_cost:,.0f}'],
                    ['ROI (First Year)', f'{roi_min:.0f}% â€“ {roi_max:.0f}%'],
                    ['Payback Period', f'{payback_max:.1f} â€“ {payback_min:.1f} months'],
                    ['3-Year Net Benefit', f'${three_year_min:,.0f} â€“ ${three_year_max:,.0f}']
                ]
            )
            doc.add_paragraph()
        
        supplier_perf = brief_data.get('supplier_performance', [])
        if supplier_perf:
            self._add_section_heading(doc, 'SUPPLIER PERFORMANCE METRICS')
            
            perf_rows = []
            for sp in supplier_perf[:15]:
                if isinstance(sp, dict):
                    # Add data source indicator
                    data_source = sp.get('data_source', 'database')
                    source_indicator = 'ðŸŒ' if data_source == 'web_search' else ''
                    
                    perf_rows.append([
                        f"{sp.get('supplier', 'N/A')} {source_indicator}".strip(),
                        f"${sp.get('spend_usd', 0):,.0f}",
                        f"{sp.get('quality_rating', 0):.1f}/5.0",
                        f"{sp.get('delivery_reliability', 0):.0f}%",
                        f"{sp.get('sustainability_score', 0):.1f}/10"
                    ])
            
            if perf_rows:
                self._create_styled_table(
                    doc,
                    ['Supplier', 'Spend (USD)', 'Quality Rating', 'Delivery %', 'Sustainability'],
                    perf_rows
                )
                
                # Add legend if any web sources were used
                has_web_sources = any(sp.get('data_source') == 'web_search' for sp in supplier_perf if isinstance(sp, dict))
                if has_web_sources:
                    legend_para = doc.add_paragraph()
                    legend_run = legend_para.add_run("ðŸŒ = Data sourced from web search (supplier not in database)")
                    legend_run.font.size = Pt(8)
                    legend_run.font.italic = True
                    
            doc.add_paragraph()
        
        rule_violations = brief_data.get('rule_violations', {})
        if rule_violations and not rule_violations.get('error'):
            self._add_section_heading(doc, 'RULE COMPLIANCE STATUS')
            
            total_v = rule_violations.get('total_violations', 0)
            total_w = rule_violations.get('total_warnings', 0)
            compliance = rule_violations.get('compliance_rate', 0)
            status = rule_violations.get('overall_status', 'UNKNOWN')
            
            self._create_styled_table(
                doc,
                ['Metric', 'Value'],
                [
                    ['Overall Status', status],
                    ['Total Violations', str(total_v)],
                    ['Total Warnings', str(total_w)],
                    ['Compliance Rate', f'{compliance:.0f}%']
                ]
            )
            
            violations_list = rule_violations.get('violations', [])
            if violations_list:
                doc.add_paragraph()
                para = doc.add_paragraph()
                run = para.add_run('Violations Detected:')
                run.font.bold = True
                
                for v in violations_list[:5]:
                    if isinstance(v, dict):
                        rule_id = v.get('rule_id', 'N/A')
                        rule_name = v.get('rule_name', 'Unknown')
                        current_val = v.get('current_value', '')
                        threshold = v.get('threshold', '')
                        
                        # Build the "why violated" explanation
                        why_violated = ""
                        if current_val and threshold:
                            why_violated = f" (Current: {current_val} vs Threshold: {threshold})"
                        
                        # Add rule with explanation
                        bullet_para = doc.add_paragraph(style='List Bullet')
                        rule_run = bullet_para.add_run(f"{rule_id}: {rule_name}")
                        rule_run.font.bold = True
                        
                        if why_violated:
                            why_run = bullet_para.add_run(why_violated)
                            why_run.font.italic = True
                            why_run.font.size = Pt(9)
            
            doc.add_paragraph()
        
        self._add_section_heading(doc, 'STRATEGIC OUTCOME')
        
        outcomes = brief_data.get('strategic_outcome', [])
        for outcome in outcomes:
            if isinstance(outcome, str):
                doc.add_paragraph(outcome, style='List Bullet')
        
        doc.add_paragraph()
        
        # Add Reasoning Sections
        why_matters = brief_data.get('why_this_matters', '')
        if why_matters:
            self._add_section_heading(doc, 'WHY THIS MATTERS')
            doc.add_paragraph(why_matters)
            doc.add_paragraph()
        
        business_just = brief_data.get('business_justification', '')
        if business_just:
            self._add_section_heading(doc, 'BUSINESS JUSTIFICATION')
            doc.add_paragraph(business_just)
            doc.add_paragraph()
        
        risk_reason = brief_data.get('risk_reasoning', '')
        if risk_reason:
            self._add_section_heading(doc, 'RISK ANALYSIS')
            doc.add_paragraph(risk_reason)
            doc.add_paragraph()
        
        rec_rationale = brief_data.get('recommendation_rationale', '')
        if rec_rationale:
            self._add_section_heading(doc, 'RECOMMENDATION RATIONALE')
            doc.add_paragraph(rec_rationale)
            doc.add_paragraph()

        # AI-GENERATED SECTIONS (GPT-4 Powered)
        if brief_data.get('llm_enabled', False):
            self._add_section_heading(doc, 'ðŸ¤– AI-GENERATED EXECUTIVE SUMMARY', level=2)
            ai_summary = brief_data.get('ai_executive_summary', '')
            if ai_summary:
                para = doc.add_paragraph(ai_summary)
                para.paragraph_format.space_after = Pt(12)
            doc.add_paragraph()

            self._add_section_heading(doc, 'ðŸ¤– AI-POWERED RISK ANALYSIS', level=2)
            ai_risk = brief_data.get('ai_risk_analysis', '')
            if ai_risk:
                para = doc.add_paragraph(ai_risk)
                para.paragraph_format.space_after = Pt(12)
            doc.add_paragraph()

            self._add_section_heading(doc, 'ðŸ¤– AI-GENERATED STRATEGIC RECOMMENDATIONS', level=2)
            ai_recs = brief_data.get('ai_strategic_recommendations', '')
            if ai_recs:
                para = doc.add_paragraph(ai_recs)
                para.paragraph_format.space_after = Pt(12)
            doc.add_paragraph()

            self._add_section_heading(doc, 'ðŸ¤– AI MARKET INTELLIGENCE', level=2)
            ai_market = brief_data.get('ai_market_intelligence', '')
            if ai_market:
                para = doc.add_paragraph(ai_market)
                para.paragraph_format.space_after = Pt(12)
            doc.add_paragraph()

        timeline = brief_data.get('implementation_timeline', [])
        if timeline:
            self._add_section_heading(doc, 'IMPLEMENTATION TIMELINE')
            
            timeline_rows = []
            for phase in timeline:
                if isinstance(phase, dict):
                    timeline_rows.append([
                        phase.get('phase', ''),
                        phase.get('duration', ''),
                        ', '.join(phase.get('activities', [])[:2])
                    ])
            
            if timeline_rows:
                self._create_styled_table(
                    doc,
                    ['Phase', 'Duration', 'Key Activities'],
                    timeline_rows
                )
            doc.add_paragraph()
        
        self._add_section_heading(doc, 'NEXT STEPS')
        
        next_steps = brief_data.get('next_steps', [])
        for i, step in enumerate(next_steps, 1):
            if isinstance(step, str):
                doc.add_paragraph(f"{i}. {step}")
        
        # Get organized subfolder based on category hierarchy
        output_folder = self._get_category_folder(brief_data)

        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        if not filename:
            category_clean = (category or 'Procurement').replace(' ', '_')
            filename = f"Incumbent_Concentration_{category_clean}_{timestamp}.docx"

        filepath = output_folder / filename
        doc.save(str(filepath))

        return str(filepath)

    def export_regional_concentration_brief(
        self, 
        brief_data: Dict[str, Any],
        filename: str = None
    ) -> str:
        """Export regional concentration brief with all enhanced sections"""
        doc = Document()
        self._set_margins(doc)
        
        category = brief_data.get('category', 'Procurement')
        total_spend = float(brief_data.get('total_spend', 0))
        fiscal_year = brief_data.get('fiscal_year', datetime.now().year)
        
        self._add_title_section(
            doc,
            brief_data.get('title', f'LEADERSHIP BRIEF â€“ {category.upper()} PROCUREMENT DIVERSIFICATION'),
            brief_data.get('subtitle', 'Regional Concentration Analysis & Diversification Strategy'),
            total_spend,
            fiscal_year
        )
        
        self._add_section_heading(doc, f'ORIGINAL SPEND CONCENTRATION ({fiscal_year})')
        
        original_concentration = brief_data.get('original_concentration', [])
        total_high_conc = self._convert_numpy(brief_data.get('total_high_concentration_pct', 0))
        total_high_spend = self._convert_numpy(brief_data.get('total_high_concentration_spend', 0))
        
        conc_rows = []
        for region in original_concentration:
            if isinstance(region, dict):
                country = region.get('country', '')
                pct = self._convert_numpy(region.get('pct', 0))
                spend = self._convert_numpy(region.get('spend_usd', 0))
                if pct > 0:
                    conc_rows.append([country, f'{pct:.0f}%', f'${spend:,.0f}'])
        
        if total_high_conc > 0:
            conc_rows.append(['Total High Concentration', f'{total_high_conc:.0f}%', f'${total_high_spend:,.0f}'])
        
        if conc_rows:
            self._create_styled_table(
                doc,
                ['Country', '% of Total Spend', 'Cost (USD)'],
                conc_rows
            )
        
        concentration_note = brief_data.get('concentration_note', '')
        if concentration_note:
            para = doc.add_paragraph(concentration_note)
            para.paragraph_format.space_before = Pt(6)
        
        doc.add_paragraph()
        
        self._add_section_heading(doc, 'DIVERSIFIED TARGET STATE (NEW SPLIT)')
        
        target_allocation = brief_data.get('target_allocation', {})
        target_rows = []
        
        if isinstance(target_allocation, dict):
            for country, data in target_allocation.items():
                if isinstance(data, dict):
                    pct = self._convert_numpy(data.get('pct', 0))
                    spend = self._convert_numpy(data.get('spend_usd', 0))
                    change = data.get('change', 'N/A')
                    target_rows.append([country, f'{pct:.0f}%', f'${spend:,.0f}', str(change)])
        
        target_rows.append(['Total', '100%', f'${total_spend:,.0f}', 'Balanced cost + risk mix'])
        
        if target_rows:
            self._create_styled_table(
                doc,
                ['Country / Region', 'New % Allocation', 'New Cost (USD)', 'Change vs Original'],
                target_rows
            )
        
        doc.add_paragraph()
        
        self._add_section_heading(doc, 'REDUCTION IN SPEND SHARE')
        
        reductions = brief_data.get('reductions', [])
        reduction_rows = []
        for reduction in reductions:
            if isinstance(reduction, dict):
                country = reduction.get('country', '')
                formatted = reduction.get('formatted_reduction', '')
                if formatted:
                    reduction_rows.append([country, formatted])
        
        if reduction_rows:
            self._create_styled_table(
                doc,
                ['Country', 'Reduction in Spend Share'],
                reduction_rows
            )
        
        doc.add_paragraph()
        
        self._add_section_heading(doc, 'REGIONAL DEPENDENCY IMPROVEMENT')
        
        regional_dep = brief_data.get('regional_dependency', {})
        corridor_name = regional_dep.get('corridor_name', 'Regional Dependency')
        original_pct = self._convert_numpy(regional_dep.get('original_pct', 90))
        new_target = regional_dep.get('new_target_pct', '58â€“63%')
        reduction_range = regional_dep.get('reduction_pct', '27â€“32%')
        
        self._create_styled_table(
            doc,
            ['Metric', f'Original {fiscal_year}', 'New Target', 'Reduction'],
            [
                [corridor_name, f'{original_pct:.0f}%', str(new_target), str(reduction_range)]
            ]
        )
        
        doc.add_paragraph()
        
        self._add_section_heading(doc, 'COST ADVANTAGE DRIVERS')
        
        cost_advantages = brief_data.get('cost_advantages', [])
        cost_rows = []
        for advantage in cost_advantages:
            if isinstance(advantage, dict):
                region = advantage.get('region', '')
                driver = advantage.get('driver', '')
                min_usd = self._convert_numpy(advantage.get('min_usd', 0))
                max_usd = self._convert_numpy(advantage.get('max_usd', 0))
                cost_rows.append([region, driver, f'${min_usd:,.0f} â€“ ${max_usd:,.0f}'])
        
        if cost_rows:
            self._create_styled_table(
                doc,
                ['Region', 'Cost Benefit Driver', 'Estimated Annual Advantage (USD)'],
                cost_rows
            )
        
        doc.add_paragraph()
        
        risk_matrix = brief_data.get('risk_matrix', {})
        if risk_matrix:
            self._add_section_heading(doc, 'RISK ASSESSMENT MATRIX')
            
            self._create_styled_table(
                doc,
                ['Risk Category', 'Current Level', 'Description'],
                [
                    ['Supply Chain Risk', risk_matrix.get('supply_chain_risk', 'N/A'), 'Supplier concentration level'],
                    ['Geographic Risk', risk_matrix.get('geographic_risk', 'N/A'), 'Regional concentration level'],
                    ['Supplier Diversity', risk_matrix.get('supplier_diversity_risk', 'N/A'), 'Number of qualified suppliers'],
                    ['Overall Risk Score', f"{risk_matrix.get('risk_score', 0)}/4.0", risk_matrix.get('overall_risk', 'N/A')]
                ]
            )
            doc.add_paragraph()
        
        roi = brief_data.get('roi_projections', {})
        if roi:
            self._add_section_heading(doc, 'ROI PROJECTIONS')
            
            self._create_styled_table(
                doc,
                ['Financial Metric', 'Projected Value'],
                [
                    ['Annual Cost Savings', f"${self._convert_numpy(roi.get('annual_cost_savings_min', 0)):,.0f} â€“ ${self._convert_numpy(roi.get('annual_cost_savings_max', 0)):,.0f}"],
                    ['Risk Reduction Value', f"${self._convert_numpy(roi.get('risk_reduction_value', 0)):,.0f}"],
                    ['ROI (First Year)', f"{roi.get('roi_percentage_min', 0):.0f}% â€“ {roi.get('roi_percentage_max', 0):.0f}%"],
                    ['3-Year Net Benefit', f"${self._convert_numpy(roi.get('three_year_net_benefit_min', 0)):,.0f} â€“ ${self._convert_numpy(roi.get('three_year_net_benefit_max', 0)):,.0f}"]
                ]
            )
            doc.add_paragraph()
        
        self._add_section_heading(doc, 'STRATEGIC OUTCOME')

        outcomes = brief_data.get('strategic_outcome', [])
        for outcome in outcomes:
            if isinstance(outcome, str):
                doc.add_paragraph(outcome, style='List Bullet')

        doc.add_paragraph()

        # REASONING SECTIONS
        why_matters = brief_data.get('why_this_matters', '')
        if why_matters:
            self._add_section_heading(doc, 'WHY THIS MATTERS')
            doc.add_paragraph(why_matters)
            doc.add_paragraph()

        business_just = brief_data.get('business_justification', '')
        if business_just:
            self._add_section_heading(doc, 'BUSINESS JUSTIFICATION')
            doc.add_paragraph(business_just)
            doc.add_paragraph()

        risk_reason = brief_data.get('risk_reasoning', '')
        if risk_reason:
            self._add_section_heading(doc, 'RISK ANALYSIS')
            doc.add_paragraph(risk_reason)
            doc.add_paragraph()

        rec_rationale = brief_data.get('recommendation_rationale', '')
        if rec_rationale:
            self._add_section_heading(doc, 'RECOMMENDATION RATIONALE')
            doc.add_paragraph(rec_rationale)
            doc.add_paragraph()

        # AI-GENERATED SECTIONS (GPT-4 Powered) for regional brief
        if brief_data.get('llm_enabled', False):
            self._add_section_heading(doc, 'ðŸ¤– AI-GENERATED EXECUTIVE SUMMARY', level=2)
            ai_summary = brief_data.get('ai_executive_summary', '')
            if ai_summary:
                para = doc.add_paragraph(ai_summary)
                para.paragraph_format.space_after = Pt(12)
            doc.add_paragraph()

            self._add_section_heading(doc, 'ðŸ¤– AI-POWERED RISK ANALYSIS', level=2)
            ai_risk = brief_data.get('ai_risk_analysis', '')
            if ai_risk:
                para = doc.add_paragraph(ai_risk)
                para.paragraph_format.space_after = Pt(12)
            doc.add_paragraph()

            self._add_section_heading(doc, 'ðŸ¤– AI-GENERATED STRATEGIC RECOMMENDATIONS', level=2)
            ai_recs = brief_data.get('ai_strategic_recommendations', '')
            if ai_recs:
                para = doc.add_paragraph(ai_recs)
                para.paragraph_format.space_after = Pt(12)
            doc.add_paragraph()

            self._add_section_heading(doc, 'ðŸ¤– AI MARKET INTELLIGENCE', level=2)
            ai_market = brief_data.get('ai_market_intelligence', '')
            if ai_market:
                para = doc.add_paragraph(ai_market)
                para.paragraph_format.space_after = Pt(12)
            doc.add_paragraph()

        timeline = brief_data.get('implementation_timeline', [])
        if timeline:
            self._add_section_heading(doc, 'IMPLEMENTATION TIMELINE')

            timeline_rows = []
            for phase in timeline:
                if isinstance(phase, dict):
                    timeline_rows.append([
                        phase.get('phase', ''),
                        phase.get('duration', ''),
                        ', '.join(phase.get('activities', [])[:2])
                    ])

            if timeline_rows:
                self._create_styled_table(
                    doc,
                    ['Phase', 'Duration', 'Key Activities'],
                    timeline_rows
                )
            doc.add_paragraph()

        self._add_section_heading(doc, 'NEXT STEPS')

        next_steps = brief_data.get('next_steps', [])
        for i, step in enumerate(next_steps, 1):
            if isinstance(step, str):
                doc.add_paragraph(f"{i}. {step}")

        # Get organized subfolder based on category hierarchy
        output_folder = self._get_category_folder(brief_data)

        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        if not filename:
            category_clean = (category or 'Procurement').replace(' ', '_')
            filename = f"Regional_Concentration_{category_clean}_{timestamp}.docx"

        filepath = output_folder / filename
        doc.save(str(filepath))

        return str(filepath)

    def export_both_briefs(
        self,
        briefs: Dict[str, Dict[str, Any]],
        export_pdf: bool = False
    ) -> Dict[str, str]:
        """Export both briefs to DOCX and optionally PDF"""
        results = {}
        
        if 'incumbent_concentration_brief' in briefs:
            incumbent_path = self.export_incumbent_concentration_brief(
                briefs['incumbent_concentration_brief']
            )
            results['incumbent_docx'] = incumbent_path
            
            if export_pdf and PDF_AVAILABLE:
                pdf_path = self.export_to_pdf(
                    briefs['incumbent_concentration_brief'], 
                    'incumbent'
                )
                results['incumbent_pdf'] = pdf_path
        
        if 'regional_concentration_brief' in briefs:
            regional_path = self.export_regional_concentration_brief(
                briefs['regional_concentration_brief']
            )
            results['regional_docx'] = regional_path
            
            if export_pdf and PDF_AVAILABLE:
                pdf_path = self.export_to_pdf(
                    briefs['regional_concentration_brief'], 
                    'regional'
                )
                results['regional_pdf'] = pdf_path
        
        return results
    
    def export_to_pdf(
        self,
        brief_data: Dict[str, Any],
        brief_type: str = 'incumbent'
    ) -> Optional[str]:
        """Export brief to PDF format"""
        if not PDF_AVAILABLE:
            print("PDF export not available. Install reportlab: pip install reportlab")
            return None
        
        category = brief_data.get('category', 'Procurement')
        total_spend = float(brief_data.get('total_spend', 0))
        fiscal_year = brief_data.get('fiscal_year', datetime.now().year)
        
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        category_clean = (category or 'Procurement').replace(' ', '_')
        
        if brief_type == 'incumbent':
            filename = f"Incumbent_Concentration_{category_clean}_{timestamp}.pdf"
        else:
            filename = f"Regional_Concentration_{category_clean}_{timestamp}.pdf"
        
        filepath = self.output_dir / filename
        
        doc = SimpleDocTemplate(
            str(filepath),
            pagesize=letter,
            rightMargin=0.75*inch,
            leftMargin=0.75*inch,
            topMargin=0.75*inch,
            bottomMargin=0.75*inch
        )
        
        styles = getSampleStyleSheet()
        
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=16,
            textColor=colors.HexColor('#003366'),
            alignment=TA_CENTER,
            spaceAfter=6
        )
        
        subtitle_style = ParagraphStyle(
            'CustomSubtitle',
            parent=styles['Normal'],
            fontSize=11,
            textColor=colors.gray,
            alignment=TA_CENTER,
            spaceAfter=12
        )
        
        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading2'],
            fontSize=12,
            textColor=colors.HexColor('#003366'),
            spaceBefore=12,
            spaceAfter=6
        )
        
        normal_style = ParagraphStyle(
            'CustomNormal',
            parent=styles['Normal'],
            fontSize=10,
            spaceAfter=6
        )
        
        story = []
        
        title = brief_data.get('title', f'LEADERSHIP BRIEF â€“ {category.upper()} DIVERSIFICATION')
        story.append(Paragraph(title, title_style))
        
        subtitle = brief_data.get('subtitle', '')
        if subtitle:
            story.append(Paragraph(subtitle, subtitle_style))
        
        story.append(Paragraph(f"<b>Total Category Spend {fiscal_year}: USD {total_spend:,.0f}</b>", subtitle_style))
        story.append(Spacer(1, 12))
        
        if brief_type == 'incumbent':
            current_state = brief_data.get('current_state', {})
            story.append(Paragraph('CURRENT STATE â€“ SUPPLIER CONCENTRATION RISK', heading_style))
            
            data = [
                ['Metric', f'{fiscal_year} Original'],
                ['Supplier', current_state.get('dominant_supplier', 'N/A')],
                ['Supplier Location', current_state.get('supplier_location', 'N/A')],
                ['Spend Share', f"{self._convert_numpy(current_state.get('spend_share_pct', 0)):.0f}%"],
                ['Key Risk', current_state.get('key_risk', 'N/A')]
            ]
            
            table = Table(data, colWidths=[2.5*inch, 4*inch])
            table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#003366')),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, -1), 9),
                ('GRID', (0, 0), (-1, -1), 0.5, colors.gray),
                ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                ('LEFTPADDING', (0, 0), (-1, -1), 6),
                ('RIGHTPADDING', (0, 0), (-1, -1), 6),
            ]))
            story.append(table)
            story.append(Spacer(1, 12))
            
            risk_statement = brief_data.get('risk_statement', '')
            if risk_statement:
                story.append(Paragraph('OVERALL RISK STATEMENT', heading_style))
                story.append(Paragraph(risk_statement, normal_style))
                story.append(Spacer(1, 12))
        
        else:
            story.append(Paragraph(f'ORIGINAL SPEND CONCENTRATION ({fiscal_year})', heading_style))
            
            original_concentration = brief_data.get('original_concentration', [])
            data = [['Country', '% of Total Spend', 'Cost (USD)']]
            for region in original_concentration:
                if isinstance(region, dict):
                    pct = self._convert_numpy(region.get('pct', 0))
                    spend = self._convert_numpy(region.get('spend_usd', 0))
                    data.append([region.get('country', ''), f'{pct:.0f}%', f'${spend:,.0f}'])
            
            if len(data) > 1:
                table = Table(data, colWidths=[2.5*inch, 2*inch, 2*inch])
                table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#003366')),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (-1, -1), 9),
                    ('GRID', (0, 0), (-1, -1), 0.5, colors.gray),
                    ('ALIGN', (1, 0), (-1, -1), 'RIGHT'),
                ]))
                story.append(table)
            
            concentration_note = brief_data.get('concentration_note', '')
            if concentration_note:
                story.append(Spacer(1, 6))
                story.append(Paragraph(concentration_note, normal_style))
        
        story.append(Spacer(1, 12))
        story.append(Paragraph('STRATEGIC OUTCOME', heading_style))
        
        outcomes = brief_data.get('strategic_outcome', [])
        for outcome in outcomes:
            if isinstance(outcome, str):
                story.append(Paragraph(f"â€¢ {outcome}", normal_style))
        
        story.append(Spacer(1, 12))
        story.append(Paragraph('NEXT STEPS', heading_style))
        
        next_steps = brief_data.get('next_steps', [])
        for i, step in enumerate(next_steps, 1):
            if isinstance(step, str):
                story.append(Paragraph(f"{i}. {step}", normal_style))
        
        doc.build(story)
        
        return str(filepath)


if __name__ == "__main__":
    print("=" * 80)
    print("DOCX EXPORTER TEST")
    print("=" * 80)
    
    test_brief = {
        'title': 'LEADERSHIP BRIEF â€“ TEST PROCUREMENT DIVERSIFICATION',
        'subtitle': 'Supplier Concentration Analysis',
        'total_spend': 2000000,
        'category': 'Test Category',
        'fiscal_year': 2024,
        'current_state': {
            'dominant_supplier': 'Test Supplier',
            'supplier_location': 'Test Region',
            'spend_share_pct': 85,
            'spend_share_usd': 1700000,
            'key_risk': 'High concentration risk'
        },
        'risk_statement': 'This is a test risk statement.',
        'strategic_outcome': ['Reduce concentration', 'Add suppliers', 'Improve resilience'],
        'next_steps': ['Step 1', 'Step 2', 'Step 3']
    }
    
    exporter = DOCXExporter()
    
    docx_path = exporter.export_incumbent_concentration_brief(test_brief, 'test_brief.docx')
    print(f"DOCX exported to: {docx_path}")
    
    if PDF_AVAILABLE:
        pdf_path = exporter.export_to_pdf(test_brief, 'incumbent')
        print(f"PDF exported to: {pdf_path}")
    else:
        print("PDF export not available - install reportlab")
