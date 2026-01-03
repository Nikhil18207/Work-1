"""
Brief Verifier - LLM-based verification of generated briefs against source data

Uses Perplexity API to fact-check generated DOCX briefs against the actual
source data (spend_data.csv) to ensure accuracy and catch any discrepancies.

Features:
- Extract key metrics from generated briefs
- Compare against source CSV data
- Use Perplexity for intelligent fact-checking
- Generate detailed verification report
"""

import os
import re
import pandas as pd
from typing import Dict, Any, List, Optional
from pathlib import Path
from docx import Document
import requests
import json


class BriefVerifier:
    """
    Verifies generated briefs against source data using Perplexity API.

    Perplexity is used because:
    - Known for accuracy and factual responses
    - Good at comparing and validating data
    - Provides confidence scores
    """

    PERPLEXITY_API_URL = "https://api.perplexity.ai/chat/completions"

    def __init__(self, api_key: Optional[str] = None):
        """
        Initialize BriefVerifier with Perplexity API.

        Args:
            api_key: Perplexity API key (if None, reads from environment)
        """
        self.api_key = api_key or os.getenv('PERPLEXITY_API_KEY', '')
        self.enabled = bool(self.api_key)

        if not self.enabled:
            print("[WARN] PERPLEXITY_API_KEY not set - verification will use basic mode")
        else:
            print("[OK] Brief Verifier initialized with Perplexity API")

    def verify_brief(
        self,
        docx_path: str,
        source_df: pd.DataFrame,
        subcategory: str
    ) -> Dict[str, Any]:
        """
        Verify a generated brief against source data.

        Args:
            docx_path: Path to the generated DOCX file
            source_df: Source DataFrame (spend_data)
            subcategory: The subcategory being verified

        Returns:
            Verification report with pass/fail status and details
        """
        # Step 1: Extract data from DOCX
        docx_data = self._extract_docx_data(docx_path)

        # Step 2: Calculate expected values from source data
        expected_data = self._calculate_expected_values(source_df, subcategory)

        # Step 3: Compare and identify discrepancies
        discrepancies = self._compare_data(docx_data, expected_data)

        # Step 4: Use Perplexity for intelligent verification (if enabled)
        if self.enabled:
            llm_verification = self._verify_with_perplexity(
                docx_data, expected_data, subcategory
            )
        else:
            llm_verification = self._basic_verification(docx_data, expected_data)

        # Step 5: Generate report
        report = {
            'success': True,
            'subcategory': subcategory,
            'docx_path': docx_path,
            'verification_method': 'perplexity' if self.enabled else 'basic',
            'docx_extracted': docx_data,
            'expected_values': expected_data,
            'discrepancies': discrepancies,
            'llm_analysis': llm_verification,
            'overall_status': 'PASS' if len(discrepancies) == 0 else 'FAIL',
            'accuracy_score': self._calculate_accuracy(discrepancies, expected_data)
        }

        return report

    def _extract_docx_data(self, docx_path: str) -> Dict[str, Any]:
        """Extract key data points from DOCX file."""
        try:
            doc = Document(docx_path)
            extracted = {
                'total_spend': None,
                'num_suppliers': None,
                'dominant_supplier': None,
                'dominant_supplier_pct': None,
                'dominant_region': None,
                'dominant_region_pct': None,
                'raw_text': ''
            }

            full_text = []

            # Extract from paragraphs
            for para in doc.paragraphs:
                text = para.text.strip()
                full_text.append(text)

                # Look for total spend - use specific patterns to avoid false matches
                # Priority 1: Look for "Total Category Spend" or "Total Spend:" patterns
                if 'total' in text.lower() and 'spend' in text.lower():
                    # Most specific: "Total Category Spend: USD X" or "Total Spend: $X"
                    total_spend_patterns = [
                        r'Total\s+(?:Category\s+)?Spend[^:]*:\s*(?:USD|US\$|\$)\s*([\d,]+(?:\.\d{2})?)',
                        r'Total\s+Spend[^:]*:\s*([\d,]+(?:\.\d{2})?)\s*(?:USD|million)?',
                    ]
                    for pattern in total_spend_patterns:
                        spend_match = re.search(pattern, text, re.IGNORECASE)
                        if spend_match:
                            value = float(spend_match.group(1).replace(',', ''))
                            if value >= 1000:
                                extracted['total_spend'] = value
                                break

                # Look for number of suppliers
                if 'supplier' in text.lower():
                    num_match = re.search(r'(\d+)\s*suppliers?', text.lower())
                    if num_match:
                        extracted['num_suppliers'] = int(num_match.group(1))

                # Look for dominant supplier percentage
                if 'dominant' in text.lower() and '%' in text:
                    pct_match = re.search(r'(\d+(?:\.\d+)?)\s*%', text)
                    if pct_match:
                        extracted['dominant_supplier_pct'] = float(pct_match.group(1))

            # Extract from tables
            for table in doc.tables:
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells]

                    # Look for supplier info in tables
                    if len(cells) >= 2:
                        if 'supplier' in cells[0].lower() and cells[1]:
                            if not extracted['dominant_supplier']:
                                extracted['dominant_supplier'] = cells[1]

                        if 'spend share' in cells[0].lower() or 'share' in cells[0].lower():
                            pct_match = re.search(r'(\d+(?:\.\d+)?)', cells[1])
                            if pct_match:
                                extracted['dominant_supplier_pct'] = float(pct_match.group(1))

                        # Region info
                        if 'region' in cells[0].lower() or 'americas' in cells[0].lower() or 'apac' in cells[0].lower():
                            if '%' in cells[1] if len(cells) > 1 else '':
                                pct_match = re.search(r'(\d+(?:\.\d+)?)', cells[1])
                                if pct_match and not extracted['dominant_region_pct']:
                                    extracted['dominant_region'] = cells[0]
                                    extracted['dominant_region_pct'] = float(pct_match.group(1))

            extracted['raw_text'] = '\n'.join(full_text[:50])  # First 50 paragraphs

            return extracted

        except Exception as e:
            return {
                'error': str(e),
                'total_spend': None,
                'num_suppliers': None,
                'dominant_supplier': None,
                'dominant_supplier_pct': None
            }

    def _calculate_expected_values(
        self,
        source_df: pd.DataFrame,
        subcategory: str
    ) -> Dict[str, Any]:
        """Calculate expected values from source data."""
        # Filter for subcategory
        if 'SubCategory' in source_df.columns:
            filtered_df = source_df[source_df['SubCategory'] == subcategory]
        else:
            filtered_df = source_df

        if filtered_df.empty:
            return {'error': f'No data found for subcategory: {subcategory}'}

        # Calculate metrics
        total_spend = filtered_df['Spend_USD'].sum()
        num_suppliers = filtered_df['Supplier_ID'].nunique()

        # Supplier analysis
        supplier_spend = filtered_df.groupby(['Supplier_ID', 'Supplier_Name'])['Spend_USD'].sum()
        supplier_spend = supplier_spend.reset_index()
        supplier_spend = supplier_spend.sort_values('Spend_USD', ascending=False)

        dominant_supplier = supplier_spend.iloc[0]['Supplier_Name'] if len(supplier_spend) > 0 else 'Unknown'
        dominant_supplier_spend = supplier_spend.iloc[0]['Spend_USD'] if len(supplier_spend) > 0 else 0
        dominant_supplier_pct = (dominant_supplier_spend / total_spend * 100) if total_spend > 0 else 0

        # Region analysis
        if 'Supplier_Region' in filtered_df.columns:
            region_spend = filtered_df.groupby('Supplier_Region')['Spend_USD'].sum()
            region_spend = region_spend.sort_values(ascending=False)
            dominant_region = region_spend.index[0] if len(region_spend) > 0 else 'Unknown'
            dominant_region_pct = (region_spend.iloc[0] / total_spend * 100) if total_spend > 0 else 0
        else:
            dominant_region = 'Unknown'
            dominant_region_pct = 0

        return {
            'total_spend': round(total_spend, 2),
            'num_suppliers': num_suppliers,
            'dominant_supplier': dominant_supplier,
            'dominant_supplier_spend': round(dominant_supplier_spend, 2),
            'dominant_supplier_pct': round(dominant_supplier_pct, 1),
            'dominant_region': dominant_region,
            'dominant_region_pct': round(dominant_region_pct, 1),
            'num_transactions': len(filtered_df)
        }

    def _compare_data(
        self,
        docx_data: Dict[str, Any],
        expected_data: Dict[str, Any]
    ) -> List[Dict[str, Any]]:
        """Compare extracted DOCX data with expected values."""
        discrepancies = []

        # Compare total spend (allow 1% tolerance)
        if docx_data.get('total_spend') and expected_data.get('total_spend'):
            docx_spend = docx_data['total_spend']
            expected_spend = expected_data['total_spend']
            tolerance = expected_spend * 0.01

            if abs(docx_spend - expected_spend) > tolerance:
                discrepancies.append({
                    'field': 'Total Spend',
                    'docx_value': f"${docx_spend:,.0f}",
                    'expected_value': f"${expected_spend:,.0f}",
                    'difference': f"${abs(docx_spend - expected_spend):,.0f}",
                    'severity': 'HIGH'
                })

        # Compare number of suppliers
        if docx_data.get('num_suppliers') and expected_data.get('num_suppliers'):
            if docx_data['num_suppliers'] != expected_data['num_suppliers']:
                discrepancies.append({
                    'field': 'Number of Suppliers',
                    'docx_value': docx_data['num_suppliers'],
                    'expected_value': expected_data['num_suppliers'],
                    'difference': abs(docx_data['num_suppliers'] - expected_data['num_suppliers']),
                    'severity': 'MEDIUM'
                })

        # Compare dominant supplier percentage (allow 1% tolerance)
        if docx_data.get('dominant_supplier_pct') and expected_data.get('dominant_supplier_pct'):
            docx_pct = docx_data['dominant_supplier_pct']
            expected_pct = expected_data['dominant_supplier_pct']

            if abs(docx_pct - expected_pct) > 1:
                discrepancies.append({
                    'field': 'Dominant Supplier %',
                    'docx_value': f"{docx_pct:.1f}%",
                    'expected_value': f"{expected_pct:.1f}%",
                    'difference': f"{abs(docx_pct - expected_pct):.1f}%",
                    'severity': 'MEDIUM'
                })

        return discrepancies

    def _verify_with_perplexity(
        self,
        docx_data: Dict[str, Any],
        expected_data: Dict[str, Any],
        subcategory: str
    ) -> Dict[str, Any]:
        """Use Perplexity API for intelligent verification."""
        try:
            prompt = f"""You are a data verification assistant. Compare the following data from a generated procurement brief against the expected source data values.

GENERATED BRIEF DATA (from DOCX):
- Total Spend: ${docx_data.get('total_spend', 'N/A'):,.0f}
- Number of Suppliers: {docx_data.get('num_suppliers', 'N/A')}
- Dominant Supplier: {docx_data.get('dominant_supplier', 'N/A')}
- Dominant Supplier Share: {docx_data.get('dominant_supplier_pct', 'N/A')}%

EXPECTED SOURCE DATA (from CSV):
- Total Spend: ${expected_data.get('total_spend', 0):,.0f}
- Number of Suppliers: {expected_data.get('num_suppliers', 'N/A')}
- Dominant Supplier: {expected_data.get('dominant_supplier', 'N/A')}
- Dominant Supplier Share: {expected_data.get('dominant_supplier_pct', 0):.1f}%
- Dominant Region: {expected_data.get('dominant_region', 'N/A')} ({expected_data.get('dominant_region_pct', 0):.1f}%)
- Number of Transactions: {expected_data.get('num_transactions', 'N/A')}

Category: {subcategory}

Please verify:
1. Do the numbers match within acceptable tolerance (1%)?
2. Are there any significant discrepancies?
3. Is the brief data accurate based on the source?

Provide a brief verification summary with:
- PASS or FAIL status
- Confidence score (0-100%)
- Any specific issues found
- Recommendation"""

            headers = {
                'Authorization': f'Bearer {self.api_key}',
                'Content-Type': 'application/json'
            }

            payload = {
                'model': 'llama-3.1-sonar-small-128k-online',
                'messages': [
                    {
                        'role': 'system',
                        'content': 'You are a precise data verification assistant. Be concise and factual.'
                    },
                    {
                        'role': 'user',
                        'content': prompt
                    }
                ],
                'temperature': 0.1,
                'max_tokens': 500
            }

            response = requests.post(
                self.PERPLEXITY_API_URL,
                headers=headers,
                json=payload,
                timeout=30
            )
            response.raise_for_status()

            result = response.json()
            analysis = result.get('choices', [{}])[0].get('message', {}).get('content', '')

            # Parse status from response
            status = 'PASS' if 'pass' in analysis.lower() and 'fail' not in analysis.lower() else 'NEEDS_REVIEW'

            return {
                'analysis': analysis,
                'status': status,
                'model': 'perplexity-sonar',
                'success': True
            }

        except Exception as e:
            return {
                'analysis': f'Perplexity verification failed: {str(e)}',
                'status': 'ERROR',
                'model': 'perplexity-sonar',
                'success': False,
                'error': str(e)
            }

    def _basic_verification(
        self,
        docx_data: Dict[str, Any],
        expected_data: Dict[str, Any]
    ) -> Dict[str, Any]:
        """Basic verification without LLM."""
        checks = []
        passed = 0
        total = 0

        # Check total spend
        if docx_data.get('total_spend') and expected_data.get('total_spend'):
            total += 1
            tolerance = expected_data['total_spend'] * 0.01
            if abs(docx_data['total_spend'] - expected_data['total_spend']) <= tolerance:
                passed += 1
                checks.append("✓ Total spend matches")
            else:
                checks.append("✗ Total spend mismatch")

        # Check supplier count
        if docx_data.get('num_suppliers') and expected_data.get('num_suppliers'):
            total += 1
            if docx_data['num_suppliers'] == expected_data['num_suppliers']:
                passed += 1
                checks.append("✓ Supplier count matches")
            else:
                checks.append("✗ Supplier count mismatch")

        # Check dominant supplier percentage
        if docx_data.get('dominant_supplier_pct') and expected_data.get('dominant_supplier_pct'):
            total += 1
            if abs(docx_data['dominant_supplier_pct'] - expected_data['dominant_supplier_pct']) <= 1:
                passed += 1
                checks.append("✓ Dominant supplier % matches")
            else:
                checks.append("✗ Dominant supplier % mismatch")

        accuracy = (passed / total * 100) if total > 0 else 0

        return {
            'analysis': '\n'.join(checks),
            'status': 'PASS' if accuracy >= 90 else 'FAIL',
            'model': 'basic',
            'success': True,
            'checks_passed': passed,
            'checks_total': total,
            'accuracy': accuracy
        }

    def _calculate_accuracy(
        self,
        discrepancies: List[Dict],
        expected_data: Dict[str, Any]
    ) -> float:
        """Calculate overall accuracy score."""
        total_fields = len([k for k in expected_data if k not in ['error', 'num_transactions']])
        errors = len(discrepancies)

        if total_fields == 0:
            return 0.0

        accuracy = ((total_fields - errors) / total_fields) * 100
        return round(accuracy, 1)

    def verify_both_briefs(
        self,
        incumbent_path: str,
        regional_path: str,
        source_df: pd.DataFrame,
        subcategory: str
    ) -> Dict[str, Any]:
        """Verify both incumbent and regional briefs."""
        results = {
            'subcategory': subcategory,
            'verification_method': 'perplexity' if self.enabled else 'basic'
        }

        # Verify incumbent brief
        if incumbent_path and Path(incumbent_path).exists():
            results['incumbent'] = self.verify_brief(incumbent_path, source_df, subcategory)
        else:
            results['incumbent'] = {'error': 'File not found', 'overall_status': 'ERROR'}

        # Verify regional brief
        if regional_path and Path(regional_path).exists():
            results['regional'] = self.verify_brief(regional_path, source_df, subcategory)
        else:
            results['regional'] = {'error': 'File not found', 'overall_status': 'ERROR'}

        # Overall status
        incumbent_pass = results['incumbent'].get('overall_status') == 'PASS'
        regional_pass = results['regional'].get('overall_status') == 'PASS'

        results['overall_status'] = 'PASS' if (incumbent_pass and regional_pass) else 'NEEDS_REVIEW'
        results['summary'] = {
            'incumbent_status': results['incumbent'].get('overall_status', 'ERROR'),
            'regional_status': results['regional'].get('overall_status', 'ERROR'),
            'incumbent_accuracy': results['incumbent'].get('accuracy_score', 0),
            'regional_accuracy': results['regional'].get('accuracy_score', 0)
        }

        return results


# Convenience function
def verify_brief(docx_path: str, source_df: pd.DataFrame, subcategory: str) -> Dict[str, Any]:
    """Quick verification of a single brief."""
    verifier = BriefVerifier()
    return verifier.verify_brief(docx_path, source_df, subcategory)
