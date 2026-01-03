"""
Verify Solar Panels briefs against source data
"""
import pandas as pd
from docx import Document
import sys
import io

# Fix encoding for Windows
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')

print('='*70)
print('VERIFICATION: Solar Panels Briefs vs Source Data')
print('='*70)

# 1. Load source data
print('\n[1] SOURCE DATA FROM spend_data.csv')
print('-'*70)
spend_df = pd.read_csv('data/structured/spend_data.csv')
solar_df = spend_df[spend_df['SubCategory'] == 'Solar Panels']

total_spend = solar_df["Spend_USD"].sum()
num_suppliers = solar_df["Supplier_ID"].nunique()
num_transactions = len(solar_df)

print(f'Total Transactions: {num_transactions}')
print(f'Total Spend: ${total_spend:,.2f}')
print(f'Unique Suppliers: {num_suppliers}')

# Supplier breakdown
print('\nSUPPLIER BREAKDOWN:')
supplier_spend = solar_df.groupby(['Supplier_ID', 'Supplier_Name', 'Supplier_Country', 'Supplier_Region'])['Spend_USD'].sum().reset_index()
supplier_spend = supplier_spend.sort_values('Spend_USD', ascending=False)

for _, row in supplier_spend.iterrows():
    pct = (row['Spend_USD'] / total_spend * 100)
    print(f"  {row['Supplier_Name']:30} | {row['Supplier_Country']:15} | {row['Supplier_Region']:12} | ${row['Spend_USD']:>12,.0f} | {pct:>5.1f}%")

# Region breakdown
print('\nREGION BREAKDOWN:')
region_spend = solar_df.groupby('Supplier_Region')['Spend_USD'].sum().sort_values(ascending=False)
for region, spend in region_spend.items():
    pct = (spend / total_spend * 100)
    print(f'  {region:15} | ${spend:>12,.0f} | {pct:>5.1f}%')

# Key metrics
dominant_supplier = supplier_spend.iloc[0]['Supplier_Name']
dominant_supplier_pct = (supplier_spend.iloc[0]['Spend_USD'] / total_spend * 100)
top3_pct = (supplier_spend.head(3)['Spend_USD'].sum() / total_spend * 100)
dominant_region = region_spend.index[0]
dominant_region_pct = (region_spend.iloc[0] / total_spend * 100)

print('\nKEY METRICS:')
print(f'  Dominant Supplier: {dominant_supplier} ({dominant_supplier_pct:.1f}%)')
print(f'  Top 3 Suppliers: {top3_pct:.1f}%')
print(f'  Dominant Region: {dominant_region} ({dominant_region_pct:.1f}%)')

# 2. Read Incumbent Brief
print('\n' + '='*70)
print('[2] INCUMBENT CONCENTRATION BRIEF - KEY VALUES')
print('-'*70)

doc = Document('outputs/briefs/Energy_and_Utilities/Solar_Panels/Incumbent_Concentration_Solar_Panels_20260104_010518.docx')

# Find key values in document
doc_values = {}
for para in doc.paragraphs:
    text = para.text.strip()
    if 'Total Category Spend' in text:
        doc_values['total_spend'] = text
    if 'dominant supplier at' in text.lower():
        doc_values['dominant_supplier'] = text[:150]
    if 'suppliers' in text.lower() and ('6' in text or 'six' in text.lower()):
        doc_values['num_suppliers'] = text[:100]

for table in doc.tables:
    for row in table.rows:
        cells = [cell.text.strip() for cell in row.cells]
        if 'Supplier' in cells[0] and len(cells) > 1:
            if cells[1] and 'Eco Power' in cells[1]:
                doc_values['table_supplier'] = f'{cells[0]}: {cells[1]}'
        if 'Spend Share' in cells[0]:
            doc_values['table_spend_share'] = f'{cells[0]}: {cells[1]}'

print('Values found in document:')
for key, value in doc_values.items():
    print(f'  {key}: {value}')

# 3. Read Regional Brief
print('\n' + '='*70)
print('[3] REGIONAL CONCENTRATION BRIEF - KEY VALUES')
print('-'*70)

doc2 = Document('outputs/briefs/Energy_and_Utilities/Solar_Panels/Regional_Concentration_Solar_Panels_20260104_010518.docx')

doc2_values = {}
for para in doc2.paragraphs:
    text = para.text.strip()
    if 'Total Category Spend' in text:
        doc2_values['total_spend'] = text
    if 'americas' in text.lower() and '%' in text:
        if len(text) < 200:
            doc2_values['americas_mention'] = text

# Get region table
for table in doc2.tables:
    for row in table.rows:
        cells = [cell.text.strip() for cell in row.cells]
        if len(cells) >= 2:
            if 'Americas' in cells[0]:
                doc2_values['americas_table'] = ' | '.join(cells)
            if 'APAC' in cells[0]:
                doc2_values['apac_table'] = ' | '.join(cells)
            if 'Middle East' in cells[0]:
                doc2_values['middle_east_table'] = ' | '.join(cells)

print('Values found in document:')
for key, value in doc2_values.items():
    print(f'  {key}: {value}')

# 4. Comparison Summary
print('\n' + '='*70)
print('[4] VERIFICATION SUMMARY')
print('='*70)

print('\n                        SOURCE DATA        DOCUMENT')
print('-'*70)
print(f'Total Spend:            ${total_spend:>12,.0f}   Check document above')
print(f'Num Suppliers:          {num_suppliers:>12}   Check document above')
print(f'Dominant Supplier:      {dominant_supplier:>12}   Check document above')
print(f'Dominant Supplier %:    {dominant_supplier_pct:>11.1f}%   Check document above')
print(f'Dominant Region:        {dominant_region:>12}   Check document above')
print(f'Dominant Region %:      {dominant_region_pct:>11.1f}%   Check document above')

print('\n' + '='*70)
print('VERIFICATION COMPLETE - Compare values above!')
print('='*70)
